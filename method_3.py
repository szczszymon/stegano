from docx import Document
from docx.shared import Pt

def define_vars():
    global CSM, group1, group2, EoS
    # Define the Character-String Mapping (CSM)
    CSM = {
        'A': 'WRZFOIN', 'B': 'F CWQXP', 'C': 'JMFXY Z', 'D': 'XA.MNTU',
        'E': 'YCHIJOS', 'F': 'KUX ZJC', 'G': ' .BCPQX', 'H': 'BLOSWMI',
        'I': 'PJGALHO', 'J': 'NSVOBGH', 'K': 'UWY.TSL', 'L': 'DIREFCQ',
        'M': 'GHTR.WB', 'N': 'ZBADEFG', 'O': 'MNSLAB.', 'P': 'TVIZM.A',
        'Q': 'SGKPHEM', 'R': 'HEUVIYJ', 'S': 'IYLNUAK', 'T': 'OFJGCDE',
        'U': 'VDEHKLY', 'V': 'LONQSUR', 'W': 'QP KXRV', 'X': 'EZDYGPT',
        'Y': '.KQUVZ ', 'Z': 'AXMTRVD', ' ': 'RTWBDNF', '.': 'CQPJ KW'
    }

    # Group definitions
    group1 = {' ', 'a', 's', 'n', 'r', 'u', 'm', 'f', 'b', 'c', 'v', 'z', 'x', 'q'}
    group2 = {'e', 't', 'o', 'i', 'h', 'd', 'l', 'w', 'g', 'y', '.', 'k', 'p', 'j'}
    EoS = ". . ."


def gather_text(doc):
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    return "\n\n".join(text), text

def shift_CSM(CSM, group1, char):
    if char in group1:
        for key in CSM.keys():
            CSM[key] = CSM[key][1::] + CSM[key][0]  # Left Circular Shift
    else:
        for key in CSM.keys():
            CSM[key] = CSM[key][-1] + CSM[key][:-1:]  # Right Circular Shift


def embed(path):
    define_vars()
    cover = Document(path)
    doc = Document()

    # Default font size detection for cover file
    def_size = cover.paragraphs[0].runs[0].font.size / 12700

    # Extraction of document data
    text, paragraphs = gather_text(cover)

    # Embedding Procedure
    secret = input("Provide a secret message: ")
    secret += EoS

    idx = 0
    x = secret[idx]
    s = CSM[x.upper()]
    txt = ""
    end = False

    # Embedding Procedure
    for paragraph in paragraphs:
        p = doc.add_paragraph()
        for y in paragraph:
            if y.upper() in s and not end:
                if txt != "":
                    r = p.add_run(txt)
                    font = r.font
                    font.size = Pt(def_size)
                    txt = ""

                pos = s.index(y.upper())
                r = p.add_run(y)
                font = r.font
                match pos:
                    case 0:
                        font.size = Pt(def_size - 1)
                    case 1:
                        font.size = Pt(def_size - 2)
                    case 2:
                        font.size = Pt(def_size - 3)
                    case 3:
                        font.size = Pt(def_size + 1)
                    case 4:
                        font.size = Pt(def_size + 2)
                    case 5:
                        font.size = Pt(def_size + 3)
                    case 6:
                        font.size = Pt(def_size + 4)

                shift_CSM(CSM, group1, x.lower())  # Perform Circular Shifts on CSM

                if idx + 1 < len(secret):
                    idx += 1
                    x = secret[idx]
                    s = CSM[x.upper()]
                else:
                    end = True

            else:
                txt += y

        if txt != "":
            r = p.add_run(txt)
            font = r.font
            font.size = Pt(def_size)
            txt = ""


    doc.save("Stego_method3.docx")


def extract(path, def_size):
    define_vars()
    doc = Document(path)

    # Default font size detection for cover file used
    for i in range(len(doc.paragraphs)):
        for j in range(len(doc.paragraphs[i].runs) - 2):
            size1 = doc.paragraphs[i].runs[j].font.size
            size2 = doc.paragraphs[i].runs[j+2].font.size

            if size1 == size2:
                def_size = size1 / 12700
                break
        else:
            continue
        break

    # Extraction Procedure
    secret = ""
    end = False

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font = run.font
            size = font.size
            if Pt(def_size) != size and not end:
                x = run.text
                if size == Pt(def_size - 1):
                    pos = 0
                elif size == Pt(def_size - 2):
                    pos = 1
                elif size == Pt(def_size - 3):
                    pos = 2
                elif size == Pt(def_size + 1):
                    pos = 3
                elif size == Pt(def_size + 2):
                    pos = 4
                elif size == Pt(def_size + 3):
                    pos = 5
                elif size == Pt(def_size + 4):
                    pos = 6

                for c in CSM.keys():
                    s = CSM[c]
                    if x.upper() not in s:
                        continue
                    if s.index(x.upper()) == pos:
                        secret += c
                        
                        shift_CSM(CSM, group1, c.lower())  # Perform Circular Shifts on CSM
                        break

                    if len(secret) >= 5:
                        if secret[-5::] == EoS:
                            end = True
                            break
                if end:
                    break
        if end:
            break

    return secret[:-5:]


embed("ref_m3.docx")
sec = extract("Stego_method3.docx", 12)
print(f"sec: {sec}")
