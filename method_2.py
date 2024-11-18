from docx import Document


# Fonts suggested by the authors of the method
default_fonts_table = {
    "Arial":                ["Arial Unicode MS", "Geo_Arial", "Microsoft Sans Serif"],
    "Book Antiqua":         ["Palatino Linotype", "Antiqua", "Caudex"],
    "Candara":              ["Khmer UI", "Ebrima", "Microsoft New Tai Lue"],
    "Century":              ["Century751 BT", "CenturyOldStyle", "CenturyExpd BT"],
    "Calibri":              ["Gisha", "Leelawadee", "Liberation Serif"],
    "Cambria":              ["Proforma", "EideticNeoRegular", "Liberation Serif"],
    "Comic Sans":           ["SF Toontime", "Komika Text", "SF Arch Rival Extended"],
    "Times New Roman":      ["Tinos", "Liberation Serif", "Thorndale"],
    "Helvetica":            ["Arimo", "Geo_Arial", "Arial-Relcom"],
    "Courier New":          ["Courier New CE", "TiredOfCourier", "TiredOfCourierThin"],
    "Verdana":              ["Tahoma", "MS Reference Sans Serif", "Lato"],
    "Perpetua":             ["ChanticleerRoman", "Centaur", "CaslonOldFace BT"],
    "Lucida Sans":          ["Lucida Sans Unicode", "Segoe UI", "Lucida Sans Typewriter"],
    "Thorndale":            ["Times New Roman", "Liberation Serif", "Tinos"],
    "Franklin Gothic Book": ["Ebrima", "Corbel", "Trebuchet MS"]
}

# Fonts available in Office 365 MS Word that have been confirmed to work
fonts_table = {
    "Candara":              ["Khmer UI", "Ebrima", "Microsoft New Tai Lue"],
    "Lucida Sans":          ["Lucida Sans Unicode", "Segoe UI", "Lucida Sans Typewriter"],
    "Franklin Gothic Book": ["Ebrima", "Corbel", "Trebuchet MS"],
    "Courier New":          ["Palatino Linotype", "Segoe UI", "Times New Roman"],
    "Aptos":                ["Courier New", "Lucida Sans", "Candara"]
}

code_table = {
    "a": (1, 1, 1), "b": (1, 1, 2), "c": (1, 1, 3),
    "d": (1, 2, 1), "e": (1, 2, 2), "f": (1, 2, 3),
    "g": (1, 3, 1), "h": (1, 3, 2), "i": (1, 3, 3),
    "j": (2, 1, 1), "k": (2, 1, 2), "l": (2, 1, 3),
    "m": (2, 2, 1), "n": (2, 2, 2), "o": (2, 2, 3),
    "p": (2, 3, 1), "q": (2, 3, 2), "r": (2, 3, 3),
    "s": (3, 1, 1), "t": (3, 1, 2), "u": (3, 1, 3),
    "v": (3, 2, 1), "w": (3, 2, 2), "x": (3, 2, 3),
    "y": (3, 3, 1), "z": (3, 3, 2), " ": (3, 3, 3)
}

reverse_code_table = dict(map(reversed, code_table.items()))

def gather_text(doc):
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    return "\n\n".join(text), text


def embed(path):
    cover = Document(path)
    doc = Document()

    # Calculation of embedding capability
    text, paragraphs = gather_text(cover)
    potential_max_len = sum(1 for c in text if c.isupper())
    max_secret_len = (potential_max_len // 3) - 3  # last 3 capital letters are used for EoS

    secret = input("Provide a secret to hide: ")

    # TODO: verify if below is needed
    if len(secret) > max_secret_len:
        return "Secret length exceeds embedding capability"

    # Finding default font in document
    ctr = 0
    while True:
        def_font = cover.paragraphs[ctr].runs[0].font.name
        ctr += 1
        if def_font is not None:
            fonts = fonts_table[def_font]
            break

    txt = ""
    idx = 0
    char_code = code_table[secret[idx]]
    ctr = -1

    # Embedding Procedure
    for paragraph in paragraphs:
        p = doc.add_paragraph()
        for char in paragraph:
            if char.isupper():
                r = p.add_run(txt)  # Saving not used text as run
                r.font.name = def_font
                txt = ""
                ctr += 1  # Rotation of character code (next font to use during embedding process)
                r = p.add_run(char)  # Encoding character from secret
                r.font.name = fonts[char_code[ctr]-1]
                if ctr == 2:  # Rotation of letters within secret
                    ctr = -1
                    idx += 1
                    if idx <= len(secret) - 1:
                        char_code = code_table[secret[idx]]
                    else:
                        break

            else:
                txt += char
        else:
            continue
        break

    doc.save("Stego_file.docx")


def extract(path):
    doc = Document(path)
    def_font = None
    fonts_used = []

    # Reading default font
    for paragraph in doc.paragraphs:
        if def_font is not None:
            break
        for run in paragraph.runs:
            if def_font is not None:
                break
            for char in run.text:
                if char.islower():
                    def_font = run.font.name
                    if def_font is not None:
                        break

    # Extraction Procedure
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for char in run.text:
                if char.isupper():
                    fonts_used.append(run.font.name)  # Saving font of all capital characters for decoding

    # Decoding Procedure
    secret = ""
    fonts = fonts_table[def_font]
    for x in range(0, len(fonts_used) - 2, 3):
        char_code = [None, None, None]
        for y in range(3):
            if fonts_used[0] == fonts_used[1] == fonts_used[2] == def_font:
                char_code = [0, 0, 0]
                break
            if fonts_used[x] == fonts[y]:
                char_code[0] = y + 1
            if fonts_used[x+1] == fonts[y]:
                char_code[1] = y + 1
            if fonts_used[x+2] == fonts[y]:
                char_code[2] = y + 1

        if tuple(char_code) == (0, 0, 0,):
            break
        else:
            secret += reverse_code_table[tuple(char_code)]

    return secret


embed("ref.docx")
sec = extract("Stego_file.docx")
print(sec)
